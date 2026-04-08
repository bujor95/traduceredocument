"""
Aplicație desktop pentru traducerea documentelor Word din română în engleză
folosind Azure Translator. Păstrează formatarea (stiluri, fonturi, tabele, headers/footers).
"""
import os
import sys
import uuid
import copy
import time
import logging
import traceback
import threading
import subprocess
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import certifi
import requests
from docx import Document
from dotenv import load_dotenv

# Asigură că requests folosește bundle-ul certifi (rezolvă SSL: CERTIFICATE_VERIFY_FAILED)
os.environ.setdefault("SSL_CERT_FILE", certifi.where())
os.environ.setdefault("REQUESTS_CA_BUNDLE", certifi.where())

def app_dir() -> str:
    """Directorul aplicației, indiferent dacă rulează ca .py sau ca .exe (PyInstaller)."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


APP_DIR = app_dir()
OUTPUT_DIR = os.path.join(APP_DIR, "rapoarte_traduse")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Caută .env atât lângă script cât și lângă .exe
load_dotenv(os.path.join(APP_DIR, ".env"))

# Logging atât în consolă cât și în fișier (translator.log lângă script).
LOG_PATH = os.path.join(APP_DIR, "translator.log")
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.FileHandler(LOG_PATH, encoding="utf-8"), logging.StreamHandler()],
)
log = logging.getLogger("translator")

AZURE_KEY = os.getenv("AZURE_TRANSLATOR_KEY", "")
AZURE_REGION = os.getenv("AZURE_TRANSLATOR_REGION", "")
AZURE_ENDPOINT = os.getenv(
    "AZURE_TRANSLATOR_ENDPOINT", "https://api.cognitive.microsofttranslator.com"
)


class AzureTranslator:
    def __init__(self, key: str, region: str, endpoint: str):
        self.key = key
        self.region = region
        self.url = endpoint.rstrip("/") + "/translate"

    def translate_batch(self, texts, from_lang="ro", to_lang="en"):
        """Traduce o listă de texte. Returnează listă de aceeași lungime."""
        if not texts:
            return []
        # Azure permite max 1000 elemente / 50.000 caractere per request.
        results = [""] * len(texts)
        batch, batch_idx, batch_chars = [], [], 0
        MAX_ITEMS, MAX_CHARS = 100, 45000

        def flush():
            nonlocal batch, batch_idx, batch_chars
            if not batch:
                return
            translated = self._call(batch, from_lang, to_lang)
            for i, t in zip(batch_idx, translated):
                results[i] = t
            batch, batch_idx, batch_chars = [], [], 0

        for i, t in enumerate(texts):
            if not t.strip():
                results[i] = t
                continue
            if len(batch) >= MAX_ITEMS or batch_chars + len(t) > MAX_CHARS:
                flush()
            batch.append(t)
            batch_idx.append(i)
            batch_chars += len(t)
        flush()
        return results

    def _call(self, texts, from_lang, to_lang, max_retries=5):
        params = {"api-version": "3.0", "from": from_lang, "to": [to_lang]}
        headers = {
            "Ocp-Apim-Subscription-Key": self.key,
            "Ocp-Apim-Subscription-Region": self.region,
            "Content-Type": "application/json",
            "X-ClientTraceId": str(uuid.uuid4()),
        }
        body = [{"text": t} for t in texts]
        last_exc = None
        for attempt in range(1, max_retries + 1):
            try:
                r = requests.post(self.url, params=params, headers=headers, json=body,
                                  timeout=60, verify=certifi.where())
                if r.status_code in (429, 500, 502, 503, 504):
                    wait = int(r.headers.get("Retry-After", 0)) or min(2 ** attempt, 30)
                    logging.warning("Azure %s la încercarea %d/%d. Aștept %ds. Body: %s",
                                    r.status_code, attempt, max_retries, wait, r.text[:500])
                    time.sleep(wait)
                    continue
                r.raise_for_status()
                data = r.json()
                return [item["translations"][0]["text"] for item in data]
            except requests.HTTPError as e:
                logging.error("HTTPError %s body=%s",
                              e, getattr(e.response, "text", "")[:1000])
                last_exc = e
                break
            except requests.RequestException as e:
                logging.warning("Network error la încercarea %d/%d: %s", attempt, max_retries, e)
                last_exc = e
                time.sleep(min(2 ** attempt, 30))
        raise last_exc if last_exc else RuntimeError("Azure: eșec după retry-uri.")


def iter_paragraphs(doc):
    """Iterează paragrafele din corp, tabele (recursiv), headere și footere."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table(table)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None:
                continue
            for p in hf.paragraphs:
                yield p
            for table in hf.tables:
                yield from _iter_table(table)


def open_file(path: str) -> None:
    """Deschide fișierul cu aplicația implicită a sistemului."""
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as e:
        logging.warning("Nu am putut deschide fișierul %s: %s", path, e)


def _iter_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:
                yield from _iter_table(t)


def translate_document(input_path: str, output_path: str, translator: AzureTranslator,
                       progress_cb=None):
    """Traduce documentul păstrând formatarea la nivel de run."""
    doc = Document(input_path)

    runs = []
    for p in iter_paragraphs(doc):
        for run in p.runs:
            if run.text:
                runs.append(run)

    if progress_cb:
        progress_cb(0, len(runs))

    texts = [r.text for r in runs]
    translated = translator.translate_batch(texts, "ro", "en")
    for run, new_text in zip(runs, translated):
        run.text = new_text

    if progress_cb:
        progress_cb(len(runs), len(runs))

    doc.save(output_path)


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Traducere document Word — RO → EN (Azure)")
        self.geometry("560x320")
        self.resizable(False, False)

        pad = {"padx": 12, "pady": 6}

        ttk.Label(self, text="Document Word (.docx):").pack(anchor="w", **pad)
        row = ttk.Frame(self); row.pack(fill="x", **pad)
        self.path_var = tk.StringVar()
        ttk.Entry(row, textvariable=self.path_var).pack(side="left", fill="x", expand=True)
        ttk.Button(row, text="Răsfoiește…", command=self.browse).pack(side="left", padx=(6, 0))

        creds = ttk.LabelFrame(self, text="Credențiale Azure Translator")
        creds.pack(fill="x", **pad)
        self.key_var = tk.StringVar(value=AZURE_KEY)
        self.region_var = tk.StringVar(value=AZURE_REGION)
        ttk.Label(creds, text="Key:").grid(row=0, column=0, sticky="w", padx=6, pady=4)
        ttk.Entry(creds, textvariable=self.key_var, show="•", width=50).grid(row=0, column=1, padx=6, pady=4)
        ttk.Label(creds, text="Region:").grid(row=1, column=0, sticky="w", padx=6, pady=4)
        ttk.Entry(creds, textvariable=self.region_var, width=50).grid(row=1, column=1, padx=6, pady=4)

        self.progress = ttk.Progressbar(self, mode="determinate")
        self.progress.pack(fill="x", **pad)
        self.status = ttk.Label(self, text="Gata.")
        self.status.pack(anchor="w", **pad)

        ttk.Button(self, text="Traduce", command=self.start).pack(pady=10)

    def browse(self):
        path = filedialog.askopenfilename(
            title="Alege document Word",
            filetypes=[("Documente Word", "*.docx")],
        )
        if path:
            self.path_var.set(path)

    def start(self):
        path = self.path_var.get().strip()
        key = AZURE_KEY
        region = AZURE_REGION
        if not path or not os.path.isfile(path):
            messagebox.showerror("Eroare", "Selectează un fișier .docx valid.")
            return
        if not key or not region:
            messagebox.showerror(
                "Configurare lipsă",
                "Lipsește AZURE_TRANSLATOR_KEY / AZURE_TRANSLATOR_REGION.\n"
                f"Creează un fișier .env lângă aplicație ({APP_DIR}).",
            )
            return

        name = os.path.basename(path)
        base, ext = os.path.splitext(name)
        out_path = os.path.join(OUTPUT_DIR, f"{base}_EN{ext}")
        translator = AzureTranslator(key, region, AZURE_ENDPOINT)

        def worker():
            try:
                log.info("=" * 60)
                log.info("Start traducere: %s", path)
                log.info("Region=%s endpoint=%s", region, AZURE_ENDPOINT)
                log.info("certifi bundle: %s", certifi.where())
                self.set_status("Se traduce…")
                translate_document(path, out_path, translator, self.set_progress)
                log.info("Salvat: %s", out_path)
                self.set_status(f"Gata: {out_path}")
                open_file(out_path)
            except requests.HTTPError as e:
                body = e.response.text if e.response is not None else ""
                log.error("HTTPError: %s\nResponse body: %s", e, body)
                log.error(traceback.format_exc())
                messagebox.showerror(
                    "Eroare Azure",
                    f"{e}\n\n{body}\n\nDetalii complete în:\n{LOG_PATH}",
                )
                self.set_status("Eroare.")
            except Exception as e:
                tb = traceback.format_exc()
                log.error("Exception: %s\n%s", e, tb)
                messagebox.showerror(
                    "Eroare",
                    f"{type(e).__name__}: {e}\n\nDetalii complete în:\n{LOG_PATH}",
                )
                self.set_status("Eroare.")

        threading.Thread(target=worker, daemon=True).start()

    def set_progress(self, current, total):
        self.progress["maximum"] = max(total, 1)
        self.progress["value"] = current
        self.update_idletasks()

    def set_status(self, text):
        self.status.config(text=text)
        self.update_idletasks()


if __name__ == "__main__":
    App().mainloop()
